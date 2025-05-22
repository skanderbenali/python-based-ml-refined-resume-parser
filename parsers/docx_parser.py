"""
DOCX Parser module for extracting text from Microsoft Word documents.
"""

import logging
import docx


class DOCXParser:
    """Parser for extracting text from DOCX files."""
    
    def parse(self, file_path):
        """
        Extract text content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content as a string
        """
        try:
            logging.debug(f"Parsing DOCX file: {file_path}")
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
            
        except Exception as e:
            logging.error(f"Error parsing DOCX file {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
